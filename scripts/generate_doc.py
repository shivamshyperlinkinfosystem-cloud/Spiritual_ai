"""Convert TECHNICAL_DOCUMENT.md to a formatted Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
SAFFRON  = RGBColor(0xD4, 0xA0, 0x17)   # #D4A017
DARK     = RGBColor(0x1A, 0x1A, 0x2E)   # #1A1A2E
TEAL     = RGBColor(0x00, 0x7A, 0x7A)   # section headings
BODY     = RGBColor(0x1A, 0x1A, 0x2E)
CODE_BG  = RGBColor(0xF4, 0xF4, 0xF4)

def set_font(run, bold=False, italic=False, size=11, colour=None):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if colour:
        run.font.color.rgb = colour

def add_title():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("🕉  Spiritual AI — Technical Document")
    set_font(r, bold=True, size=22, colour=SAFFRON)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Phase 1: AI Virtual Guru  |  Version 1.0")
    set_font(r2, italic=True, size=12, colour=TEAL)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Stack: LangGraph · Groq · ChromaDB · Streamlit")
    set_font(r3, size=11, colour=RGBColor(0x55, 0x55, 0x55))

    doc.add_paragraph()

def add_h1(text):
    doc.add_paragraph()
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = SAFFRON
        run.font.size = Pt(16)
        run.font.name = "Calibri"

def add_h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = TEAL
        run.font.size = Pt(13)
        run.font.name = "Calibri"

def add_h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = DARK
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = "Calibri"

def add_body(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        set_font(run, size=11, colour=BODY)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x00, 0x44, 0x00)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    r = p.add_run(text)
    set_font(r, size=11, colour=BODY)

def add_table_from_rows(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = "Calibri"
                run.font.size = Pt(10)
        # Saffron background
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "D4A017")
        tcPr.append(shd)

    # Data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
        # Alternating row colour
        if ri % 2 == 0:
            for ci in range(len(row)):
                tc = table.rows[ri + 1].cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  "FFF8E7")
                tcPr.append(shd)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ═══════════════════════════════════════════════════════════════════════════════

add_title()

# ── 1. Overview ───────────────────────────────────────────────────────────────
add_h1("1. Project Overview")
add_body(
    "Spiritual AI is an AI-powered virtual spiritual guide that answers questions "
    "about sacred Hindu texts — the Bhagavad Gita, Yoga Sutras of Patanjali, and "
    "Ten Principal Upanishads — using Retrieval-Augmented Generation (RAG). "
    "The system supports 5 specialized Guru personas, each with a different focus, "
    "tone, and teaching style. Every answer is grounded in actual passages from the "
    "source texts with precise citations."
)
add_h3("What the user can do")
for item in [
    "Ask questions about shlokas, sutras, and Upanishadic teachings",
    "Share personal struggles (focus, anxiety, grief, duty) and receive guidance through scriptural wisdom",
    "Switch between 5 different Guru personas",
    "See which text and page/verse each answer is sourced from",
]:
    add_bullet(item)

add_h3("What the system does NOT do")
add_bullet("Answer unrelated questions (coding, weather, sports, etc.)")
add_bullet("Fabricate answers — if the text doesn't cover a point, it says so explicitly")

# ── 2. Tech Stack ─────────────────────────────────────────────────────────────
add_h1("2. Tech Stack")
add_table_from_rows(
    ["Layer", "Technology", "Purpose"],
    [
        ["LLM", "llama-3.3-70b-versatile (Groq API)", "Generates spiritual guidance answers"],
        ["Orchestration", "LangGraph 1.2", "State machine managing the 5-node pipeline"],
        ["Vector DB", "ChromaDB", "Stores and searches text embeddings"],
        ["Embeddings", "paraphrase-multilingual-MiniLM-L12-v2 (fastembed ONNX)", "Converts text to vectors — handles Sanskrit/Devanagari"],
        ["Keyword Search", "BM25 (rank_bm25)", "Catches exact Sanskrit terms embeddings miss"],
        ["PDF Loading", "pypdf", "Extracts text from sacred text PDFs"],
        ["UI", "Streamlit 1.61", "Web chat interface"],
        ["LangChain", "langchain-core, langchain-chroma, langchain-groq", "RAG plumbing and prompt management"],
    ]
)

add_h3("Why Groq?")
add_body("Groq's LPU (Language Processing Unit) hardware runs LLaMA at ~750 tokens/second. "
         "This is critical for achieving near sub-2-second response times.")

add_h3("Why fastembed (ONNX) instead of sentence-transformers (PyTorch)?")
add_body("The sentence-transformers library requires torchvision which was not available "
         "on the deployment machine. fastembed runs the same model through ONNX Runtime — "
         "no GPU, no torchvision — and is lighter to deploy.")

add_h3("Why Hybrid Search (BM25 + Vector)?")
add_bullet("Vector search is good for semantic similarity: 'what does the Gita say about peace?'")
add_bullet("BM25 keyword search is good for exact Sanskrit terms: Nishkama Karma, Pratyahara, Samadhi")
add_bullet("Combining both gives better coverage than either alone")

# ── 3. Architecture ───────────────────────────────────────────────────────────
add_h1("3. System Architecture")
add_body("The system has three main layers:")
add_bullet("Streamlit Web App — the user interface that the person sees and types into")
add_bullet("LangGraph State Machine — the brain that routes the question through 5 processing nodes")
add_bullet("Knowledge Base — ChromaDB (vector search) + chunks.json (BM25 keyword search)")

add_h3("Architecture Diagram")
add_code(
"""USER (Browser — Streamlit)
       │
       ▼
LangGraph State Machine
  ┌──────────┐   ┌──────────┐   ┌──────────┐   ┌──────────┐
  │  guard   │──▶│ rewrite  │──▶│ retrieve │──▶│ generate │
  └──────────┘   └──────────┘   └──────────┘   └──────────┘
       │                               │               │
       └──(off-topic)──▶ reject        │               │
                                       ▼               ▼
                              ┌──────────────┐  ┌───────────┐
                              │  ChromaDB    │  │ Groq API  │
                              │  Vector DB   │  │ LLaMA 70B │
                              │  + BM25 Index│  └───────────┘
                              └──────────────┘"""
)

# ── 4. LangGraph Pipeline ─────────────────────────────────────────────────────
add_h1("4. LangGraph Pipeline (Deep Dive)")
add_body(
    "LangGraph is a framework for building stateful AI applications as a directed graph. "
    "Each node is a Python function that reads from and writes to a shared State object."
)

add_h3("State Object")
add_code(
"""class SpiritualState(TypedDict):
    messages:    list[BaseMessage]  # full conversation history
    context:     str                # retrieved passages for current question
    sources:     list[str]          # citations: "Bhagavad Gita Ch.2 V.47"
    is_relevant: bool               # guard result
    query:       str                # rewritten search query"""
)

add_h3("The 5 Nodes")
add_table_from_rows(
    ["Node", "What it does", "LLM call?"],
    [
        ["guard",    "Asks LLM: RELEVANT or IRRELEVANT? Routes accordingly",         "Yes (very short)"],
        ["rewrite",  "Rewrites follow-up questions into standalone search queries",   "Yes (if multi-turn)"],
        ["retrieve", "Hybrid search (vector + BM25) + compress to top 5 passages",   "No (vector math only)"],
        ["generate", "Persona system prompt + passages + history → LLM answer",      "Yes (main answer)"],
        ["reject",   "Returns polite refusal for off-topic questions",                "No (static message)"],
    ]
)

add_h3("Why LangGraph instead of a simple function?")
add_table_from_rows(
    ["Simple Function", "LangGraph"],
    [
        ["Hard to add/remove steps", "Add/remove nodes without touching other code"],
        ["No conditional routing", "Conditional edges (relevant / irrelevant path)"],
        ["No streaming support", "stream_mode=[updates, messages] built-in"],
        ["State is manual", "State managed automatically with reducers"],
        ["Hard to debug", "Each node's input/output is traceable"],
    ]
)

# ── 5. RAG ────────────────────────────────────────────────────────────────────
add_h1("5. RAG — How It Works")
add_body(
    "RAG = Retrieval Augmented Generation. Instead of asking the LLM from memory, "
    "we first retrieve the relevant text passages and augment the LLM's prompt with them. "
    "This grounds answers in the actual texts."
)

add_h3("Step-by-step example")
add_body('Question: "What does the Gita say about doing work without attachment?"')
add_code(
"""Step 1 — EMBED THE QUESTION
  "doing work without attachment"
  → [0.23, -0.14, 0.87, ...] (384-dimensional vector)

Step 2 — VECTOR SEARCH (ChromaDB MMR)
  Find 8 chunks most similar to the question vector
  Result: Bhagavad Gita Ch.3, Ch.5, Ch.18 passages on karma yoga

Step 3 — BM25 KEYWORD SEARCH
  Find chunks with exact words: "work", "attachment", "nishkama"
  May surface different passages than vector search

Step 4 — MERGE + COMPRESS
  Combine vector + BM25 results (deduplicate)
  Filter: keep only chunks with cosine similarity ≥ 0.28
  Keep top 5 most relevant

Step 5 — GENERATE
  System prompt (persona) + passages + question → Groq LLM
  → "Chapter 3, verse 19: 'Tasmad asaktah satatam...'"

Step 6 — CITE
  Append sources: "Bhagavad Gita Ch.3 p.45\""""
)

# ── 6. Knowledge Base ─────────────────────────────────────────────────────────
add_h1("6. Knowledge Base")
add_table_from_rows(
    ["Text", "Pages", "Chunks", "Focus"],
    [
        ["Bhagavad Gita",             "170", "919",  "Karma, Dharma, Bhakti, Jnana, all yoga"],
        ["Yoga Sutras of Patanjali",  "239", "778",  "8-limbed yoga, samadhi, meditation"],
        ["Ten Principal Upanishads",  "149", "349",  "Atman, Brahman, consciousness, non-duality"],
        ["TOTAL",                     "558", "2046", ""],
    ]
)

add_h3("PDF → Vector Store pipeline")
add_code(
"""PDF File  (pypdf)
   │
   ▼
Raw text per page + metadata: {page, source_text, chapter, verse}
   │
   ▼  RecursiveCharacterTextSplitter
Chunks ~800 chars, 150-char overlap
Sanskrit-aware separators: ॥  ।  \\n\\n  \\n
   │
   ├──▶ chunks.json   (BM25 keyword search)
   │
   ▼  fastembed ONNX
384-dimensional vectors
   │
   ▼  ChromaDB
Stored in chroma_db/ with metadata"""
)

add_h3("Adding new PDFs")
add_body("Drop any .pdf file into the project folder and run: python ingest.py")
add_body("The system auto-discovers all PDFs, assigns source names, and rebuilds the knowledge base.")

# ── 7. Personas ───────────────────────────────────────────────────────────────
add_h1("7. Five Guru Personas")
add_body(
    "Each persona is a different system prompt injected at the top of every LLM call. "
    "The RAG pipeline and knowledge base are identical for all personas."
)
add_table_from_rows(
    ["Persona", "Focus", "Tone", "Primary Texts"],
    [
        ["🕉 General Guru",           "Balanced across all texts",        "Warm, accessible",         "All three"],
        ["🙏 Bhakti Guru",            "Devotion, love, surrender",        "Heart-centred, devotional", "Gita Ch.9,12,18"],
        ["🧘 Yoga Guru",              "8 limbs, practice, discipline",    "Systematic, practical",     "Yoga Sutras"],
        ["🔮 Meditation Guru",        "Consciousness, Atman, silence",    "Contemplative, still",      "Upanishads"],
        ["⚖️ Karma & Dharma Guru",   "Right action, duty, ethics",       "Direct, grounded",          "Gita Ch.2,3,5"],
        ["💚 Healing Guru",           "Grief, anxiety, emotional pain",   "Compassionate, gentle",     "Gita Ch.2,18 + Upanishads"],
    ]
)

# ── 8. File Structure ─────────────────────────────────────────────────────────
add_h1("8. File Structure")
add_code(
"""Spiritual_AI/
│
├── app.py                     ← Streamlit UI (persona selector + streaming chat)
├── geeta_chat.py              ← LangGraph pipeline + 5 personas + RAG logic
├── ingest.py                  ← PDF → ChromaDB ingestion pipeline
├── requirements.txt           ← Python dependencies
│
├── final_geeta.pdf            ← Bhagavad Gita (Sanskrit)
├── Patanjali-yogasutra.IGS.pdf   ← Yoga Sutras of Patanjali
├── The-Ten-Principal-Upanishads.pdf  ← Ten Principal Upanishads
│
├── chroma_db/                 ← ChromaDB vector store (auto-created)
│   ├── chroma.sqlite3         ← Main database (14 MB)
│   └── <uuid>/                ← HNSW index files
│
├── chunks.json                ← 2046 chunks for BM25 keyword search
│
├── .streamlit/
│   └── config.toml            ← Theme (warm dark spiritual colours)
│
└── TECHNICAL_DOCUMENT.md     ← This document"""
)

add_h3("Key file responsibilities")
add_table_from_rows(
    ["File", "Role"],
    [
        ["geeta_chat.py", "Brain — defines PERSONAS, SpiritualState, all 5 nodes, build_app()"],
        ["app.py",        "Face — Streamlit UI, persona selector, streaming chat loop"],
        ["ingest.py",     "Librarian — auto-discovers PDFs, builds ChromaDB + chunks.json"],
    ]
)

# ── 9. Data Flow ──────────────────────────────────────────────────────────────
add_h1("9. End-to-End Data Flow")
add_body('Example: User types "I feel anxious. How can I find peace?"')
add_code(
"""1. Streamlit  →  Appends HumanMessage to session_state.messages

2. guard node →  Groq prompt: "RELEVANT or IRRELEVANT?"
                 Response: "RELEVANT"
                 state.is_relevant = True

3. rewrite node → First message, no history
                  state.query = original question

4. retrieve node → Vector search: 8 chunks (Gita Ch.2, Upanishads)
                   BM25 search:   8 chunks
                   Merge + deduplicate: 12 unique chunks
                   Compress (cosine ≥ 0.28): 5 best chunks
                   state.context = [5 passages]
                   state.sources = ["Gita p.19", "Upanishads p.43"]

5. generate node → System prompt (Healing Guru persona)
                   + 5 retrieved passages
                   + conversation history
                   → Groq streams back the answer token by token

6. Streamlit    → Renders each token live with ▌ cursor
                  Shows step indicators while processing
                  Renders source citations in expander"""
)

# ── 10. API Details ───────────────────────────────────────────────────────────
add_h1("10. API & Model Details")

add_h3("Groq API")
add_table_from_rows(
    ["Parameter", "Value"],
    [
        ["Model",        "llama-3.3-70b-versatile"],
        ["Temperature",  "0.1 (low — consistent, factual answers)"],
        ["Max tokens",   "2048 per response"],
        ["API key",      "Environment variable: GROQ_API_KEY"],
        ["Speed",        "~750 tokens/second (LPU hardware)"],
    ]
)

add_h3("Embedding Model")
add_table_from_rows(
    ["Parameter", "Value"],
    [
        ["Model",         "paraphrase-multilingual-MiniLM-L12-v2"],
        ["Provider",      "fastembed (ONNX Runtime — no GPU needed)"],
        ["Dimensions",    "384"],
        ["Languages",     "50+ including Hindi/Sanskrit (Devanagari)"],
        ["Download size", "~45 MB (cached after first run)"],
    ]
)

add_h3("ChromaDB Settings")
add_table_from_rows(
    ["Parameter", "Value"],
    [
        ["Search type",  "MMR — Maximal Marginal Relevance (diverse results)"],
        ["Fetch k",      "20 candidates"],
        ["Return k",     "8 results"],
        ["After compress", "Top 5 by cosine similarity"],
        ["Similarity threshold", "≥ 0.28 to keep a chunk"],
    ]
)

# ── 11. Streamlit UI ──────────────────────────────────────────────────────────
add_h1("11. Streamlit UI")
add_h3("Caching Strategy")
add_table_from_rows(
    ["Cache", "Function", "Reloaded when"],
    [
        ["@st.cache_resource",               "load_base() — embeddings + ChromaDB", "App restarts"],
        ["@st.cache_resource (per persona)", "load_persona_app(key)",               "New persona first use"],
        ["st.session_state",                 "messages, chat_display",              "Page refresh or Clear"],
    ]
)
add_body("Each persona's LangGraph graph is compiled once and cached. Switching personas is instant after first use.")

add_h3("Streaming")
add_body("The app uses LangGraph's dual stream mode to show live step indicators AND stream the answer:")
add_code(
"""for mode, data in app.stream(state, stream_mode=["updates", "messages"]):
    if mode == "updates":
        # Node completed → update step indicator
        node = next(iter(data))
        step_placeholder.caption("📖 Searching texts...")

    elif mode == "messages":
        # LLM token received → stream to UI
        chunk, metadata = data
        if metadata["langgraph_node"] == "generate":
            full_content += chunk.content
            answer_placeholder.markdown(full_content + "▌")"""
)

# ── 12. Setup ─────────────────────────────────────────────────────────────────
add_h1("12. Setup & Run")
add_code(
"""# 1. Clone the repository
git clone https://github.com/shivamshyperlinkinfosystem-cloud/Spiritual_ai.git
cd Spiritual_ai

# 2. Create virtual environment
python -m venv venv
source venv/bin/activate    # Linux/Mac
# OR: venv\\Scripts\\activate  (Windows)

# 3. Install dependencies
pip install -r requirements.txt

# 4. Build the knowledge base (one-time)
python ingest.py

# 5. Set API key and run
export GROQ_API_KEY="gsk_your_key_here"
streamlit run app.py"""
)

# ── 13. Deployment ────────────────────────────────────────────────────────────
add_h1("13. Deployment on Streamlit Cloud")
add_h3("Steps")
for step in [
    "Go to share.streamlit.io",
    "Click New app → connect GitHub repo: shivamshyperlinkinfosystem-cloud/Spiritual_ai",
    "Branch: main  |  Main file: app.py",
    "Click Advanced settings → Secrets and add: GROQ_API_KEY = \"gsk_your_key\"",
    "Click Deploy — no ingestion needed (chroma_db/ is in the repo)",
]:
    add_bullet(step)

add_table_from_rows(
    ["Resource", "Free Tier", "Our Usage"],
    [
        ["RAM",  "1 GB",   "~400 MB (embeddings + ChromaDB)"],
        ["CPU",  "Shared", "ONNX embedding — CPU-only, no GPU needed"],
        ["Disk", "1 GB",   "~50 MB (chroma_db + chunks.json)"],
    ]
)

# ── 14. Phase 1 Status ────────────────────────────────────────────────────────
add_h1("14. Phase 1 Completion Status")
add_table_from_rows(
    ["Requirement", "Status", "Notes"],
    [
        ["RAG infrastructure — vector DB + embedding pipeline", "✅ Done", "ChromaDB + fastembed ONNX"],
        ["RAG infrastructure — ingestion tooling", "✅ Done", "ingest.py auto-discovers PDFs"],
        ["AI Virtual Guru — 5 specialized personas", "✅ Done", "Bhakti, Yoga, Meditation, Karma, Healing"],
        ["Knowledge base — Bhagavad Gita", "✅ Done", "919 chunks"],
        ["Knowledge base — Upanishads", "✅ Done", "349 chunks (Ten Principal)"],
        ["Knowledge base — foundational texts", "✅ Done", "Yoga Sutras — 778 chunks"],
        ["Multi-turn conversational interface", "✅ Done", "Full history in LangGraph state"],
        ["Source grounding with citations", "✅ Done", "Ch/Verse + source text name"],
        ["Topic guard — reject off-topic questions", "✅ Done", "guard node with LLM classification"],
        ["Streaming responses", "✅ Done", "stream_mode=[updates, messages]"],
        ["Hybrid search (vector + BM25)", "✅ Done", "Combined retrieval"],
        ["Query rewriting for follow-ups", "✅ Done", "rewrite node"],
        ["Streamlit UI with persona selector", "✅ Done", "Sidebar radio buttons"],
        ["GitHub repository", "✅ Done", "Clean — source files only"],
        ["Deployment ready", "✅ Done", "Streamlit Cloud compatible"],
        ["Sub-2-second response", "⚠️ Partial", "Answer streams in ~1-2s; full pipeline ~4-5s"],
    ]
)

# ── 15. Phase 2 ───────────────────────────────────────────────────────────────
add_h1("15. What to Build Next (Phase 2)")

add_h3("1. Reduce response time (biggest UX improvement)")
for item in [
    "Run guard + rewrite in parallel — cut 1.5 seconds",
    "Use a smaller/faster model for guard + rewrite (e.g., Haiku 4.5)",
    "Skip rewrite for first message (already done for first turn)",
    "Cache guard results for repeated questions",
]:
    add_bullet(item)

add_h3("2. More sacred texts (easy to add)")
for item in [
    "Narada Bhakti Sutras — for Bhakti persona",
    "Vivekachudamani by Shankaracharya — for Meditation persona",
    "Hatha Yoga Pradipika — for Yoga persona",
    "Devi Mahatmyam — for Healing persona",
]:
    add_bullet(item)

add_h3("3. Voice interface")
add_bullet("User speaks → Whisper API transcribes → Guru answers → Text-to-speech reads aloud")

add_h3("4. User accounts + conversation history")
add_bullet("Save conversations per user")
add_bullet("Continue from last session")
add_bullet("Personal spiritual journal")

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "Spiritual_AI_Technical_Document.docx"
doc.save(output_path)
print(f"✅  Word document saved: {output_path}")
